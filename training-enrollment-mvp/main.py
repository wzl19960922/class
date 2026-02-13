import hashlib
import io
import json
import csv
import os
import re
import sqlite3
import zipfile
import base64
import logging
import urllib.error
import urllib.parse
import urllib.request
from logging.handlers import RotatingFileHandler
from datetime import date, datetime, timedelta
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

import pandas as pd
import qrcode
from docx import Document
from flask import Flask, jsonify, render_template, request, send_file

try:
    import PIL  # noqa: F401
    QR_PIL_AVAILABLE = True
except Exception:
    QR_PIL_AVAILABLE = False

BASE_DIR = Path(__file__).resolve().parent
DB_PATH = BASE_DIR / "training.db"
UPLOAD_DIR = BASE_DIR / "uploads"
LOG_DIR = BASE_DIR / "logs"
LOG_PATH = LOG_DIR / "app.log"

app = Flask(__name__)

LATEST_SESSION_ID: Optional[int] = None


def setup_logging() -> None:
    LOG_DIR.mkdir(exist_ok=True)
    handler = RotatingFileHandler(LOG_PATH, maxBytes=1_000_000, backupCount=3, encoding="utf-8")
    handler.setFormatter(
        logging.Formatter("%(asctime)s [%(levelname)s] %(message)s")
    )
    app.logger.setLevel(logging.INFO)
    app.logger.handlers.clear()
    app.logger.addHandler(handler)
    app.logger.addHandler(logging.StreamHandler())


ALLOWED_EXCEL_EXTENSIONS = {".xlsx", ".xlsm", ".xltx", ".xltm", ".xls"}
ALLOWED_WORD_EXTENSIONS = {".docx"}


def is_excel_filename(filename: str) -> bool:
    return Path(filename).suffix.lower() in ALLOWED_EXCEL_EXTENSIONS


def is_word_filename(filename: str) -> bool:
    return Path(filename).suffix.lower() in ALLOWED_WORD_EXTENSIONS


def get_connection() -> sqlite3.Connection:
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    return conn


def initialize_database() -> None:
    with get_connection() as conn:
        conn.execute(
            """
            CREATE TABLE IF NOT EXISTS person (
                person_id INTEGER PRIMARY KEY AUTOINCREMENT,
                phone_norm TEXT UNIQUE NOT NULL,
                name_latest TEXT,
                org_text_latest TEXT
            )
            """
        )
        conn.execute(
            """
            CREATE TABLE IF NOT EXISTS training_session (
                session_id INTEGER PRIMARY KEY AUTOINCREMENT,
                title TEXT,
                start_date TEXT,
                end_date TEXT,
                location_text TEXT,
                training_goal TEXT,
                notice_filename TEXT,
                notice_sha256 TEXT,
                created_at TEXT
            )
            """
        )
        columns = {
            row["name"] for row in conn.execute("PRAGMA table_info(training_session)").fetchall()
        }
        if "training_goal" not in columns:
            conn.execute("ALTER TABLE training_session ADD COLUMN training_goal TEXT")
        conn.execute(
            """
            CREATE TABLE IF NOT EXISTS enrollment (
                enrollment_id INTEGER PRIMARY KEY AUTOINCREMENT,
                session_id INTEGER NOT NULL,
                person_id INTEGER NOT NULL,
                enrolled_at TEXT,
                name_snapshot TEXT,
                org_text TEXT,
                region_text TEXT,
                title_text TEXT,
                remote_id_snapshot TEXT,
                room_preference TEXT,
                source_file TEXT,
                source_sheet TEXT,
                FOREIGN KEY(session_id) REFERENCES training_session(session_id),
                FOREIGN KEY(person_id) REFERENCES person(person_id)
            )
            """
        )
        conn.execute(
            """
            CREATE TABLE IF NOT EXISTS course (
                course_id INTEGER PRIMARY KEY AUTOINCREMENT,
                title TEXT NOT NULL,
                teacher TEXT,
                start_at TEXT,
                end_at TEXT,
                location TEXT,
                session_id INTEGER,
                source_file TEXT,
                created_at TEXT,
                FOREIGN KEY(session_id) REFERENCES training_session(session_id)
            )
            """
        )
        conn.execute(
            """
            CREATE TABLE IF NOT EXISTS message_task (
                task_id INTEGER PRIMARY KEY AUTOINCREMENT,
                course_id INTEGER NOT NULL,
                task_type TEXT NOT NULL,
                planned_at TEXT NOT NULL,
                content TEXT,
                survey_link TEXT,
                qr_data_uri TEXT,
                status TEXT NOT NULL DEFAULT 'pending',
                sent_at TEXT,
                created_at TEXT,
                UNIQUE(course_id, task_type, planned_at),
                FOREIGN KEY(course_id) REFERENCES course(course_id)
            )
            """
        )
        conn.execute(
            """
            CREATE TABLE IF NOT EXISTS survey_response (
                response_id INTEGER PRIMARY KEY AUTOINCREMENT,
                course_id INTEGER NOT NULL,
                satisfaction_score INTEGER,
                gain_text TEXT,
                suggestion_text TEXT,
                recommend_score INTEGER,
                submitted_at TEXT,
                FOREIGN KEY(course_id) REFERENCES course(course_id)
            )
            """
        )
        conn.execute(
            """
            CREATE TABLE IF NOT EXISTS finance_record (
                record_id INTEGER PRIMARY KEY AUTOINCREMENT,
                record_no TEXT UNIQUE,
                start_time TEXT,
                end_time TEXT,
                duration_text TEXT,
                name TEXT,
                phone TEXT,
                id_card TEXT,
                org_name TEXT,
                job_title TEXT,
                bank_card TEXT,
                bank_name TEXT,
                city_name TEXT,
                user_type TEXT,
                nickname TEXT,
                source_file TEXT,
                updated_at TEXT,
                raw_json TEXT
            )
            """
        )


def normalize_phone(value: Any) -> Optional[str]:
    if value is None:
        return None
    text = str(value).strip()
    if not text:
        return None
    text = text.replace(" ", "").replace("-", "")
    if text.startswith("+86"):
        text = text[3:]
    if text.startswith("86"):
        text = text[2:]
    digits = re.sub(r"\D", "", text)
    if len(digits) >= 11:
        digits = digits[-11:]
    if re.fullmatch(r"1\d{10}", digits):
        return digits
    return None


def guess_column(columns: List[str], keywords: List[str]) -> Optional[str]:
    normalized = {col: re.sub(r"\s+", "", col).lower() for col in columns}
    for col, col_norm in normalized.items():
        for keyword in keywords:
            if keyword in col_norm:
                return col
    return None


def row_has_data(values: List[Any]) -> bool:
    for value in values:
        if value is None:
            continue
        if isinstance(value, float) and pd.isna(value):
            continue
        if str(value).strip() != "":
            return True
    return False


def save_upload(file_storage) -> Tuple[str, str]:
    UPLOAD_DIR.mkdir(exist_ok=True)
    filename = file_storage.filename or "upload"
    timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
    safe_name = re.sub(r"[^A-Za-z0-9_.-]", "_", filename)
    saved_name = f"{timestamp}_{safe_name}"
    file_path = UPLOAD_DIR / saved_name
    file_storage.save(file_path)
    return saved_name, str(file_path)


def compute_sha256(file_path: str) -> str:
    sha256 = hashlib.sha256()
    with open(file_path, "rb") as handle:
        for chunk in iter(lambda: handle.read(8192), b""):
            sha256.update(chunk)
    return sha256.hexdigest()


def json_response(ok: bool, data: Any = None, error: Optional[str] = None):
    return jsonify({"ok": ok, "data": data, "error": error})


@app.errorhandler(Exception)
def handle_exception(exc: Exception):
    app.logger.exception("Unhandled exception on %s %s", request.method, request.path)
    if request.path.startswith("/api/"):
        return json_response(False, error=f"服务异常: {exc}"), 500
    raise exc


@app.route("/", endpoint="home_page")
def index():
    return render_template("index.html")


@app.route("/add")
def add_page():
    return render_template("add.html")


@app.route("/api/session/create", methods=["POST"])
def create_session():
    title = request.form.get("title", "").strip()
    start_date = request.form.get("start_date", "").strip()
    end_date = request.form.get("end_date", "").strip()
    location_text = request.form.get("location_text", "").strip()
    training_goal = request.form.get("training_goal", "").strip()

    notice_file = request.files.get("notice_file")
    notice_filename = None
    notice_sha256 = None
    if notice_file and notice_file.filename:
        notice_filename, file_path = save_upload(notice_file)
        notice_sha256 = compute_sha256(file_path)

    created_at = datetime.now().isoformat(timespec="seconds")
    with get_connection() as conn:
        cursor = conn.execute(
            """
            INSERT INTO training_session
            (title, start_date, end_date, location_text, training_goal, notice_filename, notice_sha256, created_at)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                title,
                start_date,
                end_date,
                location_text,
                training_goal,
                notice_filename,
                notice_sha256,
                created_at,
            ),
        )
        session_id = cursor.lastrowid

    global LATEST_SESSION_ID
    LATEST_SESSION_ID = session_id
    return json_response(True, {"session_id": session_id})


@app.route("/api/session/<int:session_id>")
def get_session(session_id: int):
    with get_connection() as conn:
        row = conn.execute(
            """
            SELECT session_id, title, start_date, end_date, location_text,
                   training_goal, notice_filename, notice_sha256, created_at
            FROM training_session
            WHERE session_id = ?
            """,
            (session_id,),
        ).fetchone()
    if not row:
        return json_response(False, error="期次不存在。")
    return json_response(True, dict(row))


@app.route("/api/session/update", methods=["POST"])
def update_session():
    session_id_text = request.form.get("session_id", "").strip()
    if not session_id_text.isdigit():
        return json_response(False, error="session_id 非法。")

    session_id = int(session_id_text)
    title = request.form.get("title", "").strip()
    start_date = request.form.get("start_date", "").strip()
    end_date = request.form.get("end_date", "").strip()
    location_text = request.form.get("location_text", "").strip()
    training_goal = request.form.get("training_goal", "").strip()

    with get_connection() as conn:
        existing = conn.execute(
            """
            SELECT notice_filename, notice_sha256
            FROM training_session
            WHERE session_id = ?
            """,
            (session_id,),
        ).fetchone()
        if not existing:
            return json_response(False, error="期次不存在。")

        notice_filename = existing["notice_filename"]
        notice_sha256 = existing["notice_sha256"]
        notice_file = request.files.get("notice_file")
        if notice_file and notice_file.filename:
            notice_filename, file_path = save_upload(notice_file)
            notice_sha256 = compute_sha256(file_path)

        conn.execute(
            """
            UPDATE training_session
            SET title = ?, start_date = ?, end_date = ?, location_text = ?, training_goal = ?,
                notice_filename = ?, notice_sha256 = ?
            WHERE session_id = ?
            """,
            (
                title,
                start_date,
                end_date,
                location_text,
                training_goal,
                notice_filename,
                notice_sha256,
                session_id,
            ),
        )

    global LATEST_SESSION_ID
    LATEST_SESSION_ID = session_id
    return json_response(True, {"session_id": session_id})


def resolve_session_id(session_id_value: Optional[str]) -> Optional[int]:
    if session_id_value:
        try:
            return int(session_id_value)
        except ValueError:
            return None
    return LATEST_SESSION_ID


def extract_notice_text(file_path: str) -> str:
    suffix = Path(file_path).suffix.lower()
    if suffix == ".docx":
        doc = Document(file_path)
        lines: List[str] = []
        for p in doc.paragraphs:
            text = (p.text or "").strip()
            if text:
                lines.append(text)
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join((cell.text or "").strip() for cell in row.cells if (cell.text or "").strip())
                if row_text:
                    lines.append(row_text)
        return "\n".join(lines)

    with open(file_path, "r", encoding="utf-8", errors="ignore") as handle:
        return handle.read()


def parse_json_from_text(text: str) -> Dict[str, str]:
    cleaned = (text or "").strip()
    if cleaned.startswith("```"):
        cleaned = re.sub(r"^```(?:json)?", "", cleaned).strip()
        cleaned = re.sub(r"```$", "", cleaned).strip()
    parsed = json.loads(cleaned)
    if not isinstance(parsed, dict):
        raise ValueError("模型返回格式不是 JSON 对象")
    return {
        "title": str(parsed.get("title", "")).strip(),
        "start_date": str(parsed.get("start_date", "")).strip(),
        "end_date": str(parsed.get("end_date", "")).strip(),
        "location_text": str(parsed.get("location_text", "")).strip(),
        "training_goal": str(parsed.get("training_goal", "")).strip(),
    }


def parse_notice_with_baidu_llm(notice_text: str, api_key: str) -> Dict[str, str]:
    endpoint = "https://qianfan.baidubce.com/v2/chat/completions"
    prompt = (
        "你是信息抽取助手。请从以下培训通知文本提取字段，并且只输出 JSON，不要输出其它内容。"
        "\n字段：title(培训班名称),start_date(YYYY-MM-DD),end_date(YYYY-MM-DD),location_text(培训地点),training_goal(培训目标)。"
        "\n若某项缺失填空字符串。\n\n通知文本：\n"
        f"{notice_text[:12000]}"
    )
    body = json.dumps(
        {
            "model": "ernie-4.5-turbo-128k",
            "messages": [{"role": "user", "content": prompt}],
            "temperature": 0.01,
            "top_p": 0.8,
        }
    ).encode("utf-8")
    req = urllib.request.Request(
        endpoint,
        data=body,
        headers={
            "Content-Type": "application/json",
            "Authorization": f"Bearer {api_key}",
        },
        method="POST",
    )
    with urllib.request.urlopen(req, timeout=40) as resp:
        raw = resp.read().decode("utf-8")
    payload = json.loads(raw)
    if payload.get("error"):
        message = payload["error"].get("message") or payload["error"].get("type") or "接口返回错误"
        raise ValueError(message)

    result_text = payload.get("result", "")
    if not result_text and payload.get("choices"):
        first_choice = payload["choices"][0] if payload["choices"] else {}
        message = first_choice.get("message", {}) if isinstance(first_choice, dict) else {}
        result_text = message.get("content", "")
    return parse_json_from_text(result_text)


def build_map_info(location_text: str, amap_key: str) -> Dict[str, str]:
    location_text = (location_text or "").strip()
    if not location_text:
        return {"map_url": "", "geo": ""}

    query = urllib.parse.urlencode({"query": location_text})
    map_url = f"https://uri.amap.com/search?{query}"
    if not amap_key:
        return {"map_url": map_url, "geo": ""}

    geocode_params = urllib.parse.urlencode(
        {"address": location_text, "key": amap_key, "output": "json"}
    )
    geocode_url = f"https://restapi.amap.com/v3/geocode/geo?{geocode_params}"
    try:
        req = urllib.request.Request(geocode_url, method="GET")
        with urllib.request.urlopen(req, timeout=10) as resp:
            payload = json.loads(resp.read().decode("utf-8"))
        geocodes = payload.get("geocodes") or []
        if geocodes and isinstance(geocodes[0], dict):
            geo = geocodes[0].get("location", "")
            return {"map_url": map_url, "geo": geo}
    except Exception:
        app.logger.exception("Map geocode failed for location=%s", location_text)
    return {"map_url": map_url, "geo": ""}


@app.route("/api/session/parse_notice", methods=["POST"], endpoint="api_session_parse_notice")
def parse_notice_api():
    notice_file = request.files.get("notice_file")
    if not notice_file or not notice_file.filename:
        return json_response(False, error="请先选择通知文件。")

    suffix = Path(notice_file.filename).suffix.lower()
    if suffix not in {".docx", ".txt"}:
        return json_response(False, error="目前仅支持 .docx 或 .txt 通知文件解析。")

    api_key = request.form.get("baidu_api_key", "").strip()
    if not api_key:
        return json_response(False, error="请填写百度千帆 API Key（Bearer）。")

    _, file_path = save_upload(notice_file)
    try:
        notice_text = extract_notice_text(file_path)
        if not notice_text.strip():
            return json_response(False, error="通知文件未读取到有效文本，请检查文档内容。")
        parsed = parse_notice_with_baidu_llm(notice_text, api_key)
        return json_response(True, parsed)
    except urllib.error.HTTPError as exc:
        detail = exc.read().decode("utf-8", errors="ignore")
        app.logger.exception("Baidu parse HTTP error: %s", detail)
        return json_response(False, error=f"百度云接口调用失败：{detail[:300] or exc.reason}")
    except ValueError as exc:
        message = str(exc)
        if "Access token invalid" in message or "Invalid authentication" in message:
            return json_response(False, error="API Key 无效或已过期，请在百度千帆控制台重新获取后再试。")
        return json_response(False, error=f"解析失败：{message}")
    except Exception as exc:
        app.logger.exception("Parse notice failed")
        return json_response(False, error=f"解析失败：{exc}")


def import_excel(file_path: str, source_file: str, session_id: int) -> Dict[str, Any]:
    sheets = pd.read_excel(file_path, sheet_name=None, dtype=str, engine="openpyxl")
    sheet_count = len(sheets)
    valid_rows = 0
    new_person_count = 0
    new_enrollment_count = 0
    exceptions: List[Dict[str, Any]] = []

    with get_connection() as conn:
        try:
            for sheet_name, df in sheets.items():
                if df is None or df.empty:
                    continue
                df = df.fillna("")
                columns = list(df.columns)
                phone_col = guess_column(
                    columns, ["手机", "手机号", "电话", "mobile", "phone"]
                )
                if not phone_col:
                    exceptions.append(
                        {
                            "sheet": sheet_name,
                            "row": None,
                            "reason": "未找到手机号列",
                        }
                    )
                    continue

                name_col = guess_column(columns, ["姓名", "name"])
                org_col = guess_column(columns, ["单位", "机构", "company", "org"])
                region_col = guess_column(columns, ["地区", "区域", "省", "市", "region"])
                title_col = guess_column(columns, ["职务", "岗位", "title"])
                remote_id_col = guess_column(columns, ["工号", "编号", "学号", "id"])
                room_col = guess_column(columns, ["住宿", "房间", "room"])

                column_index = {col: idx for idx, col in enumerate(columns)}
                for row_index, row in enumerate(
                    df.itertuples(index=False, name=None), start=2
                ):
                    values = list(row)
                    if not row_has_data(values):
                        continue
                    phone_raw = row[column_index[phone_col]]
                    phone_norm = normalize_phone(phone_raw)
                    if not phone_norm:
                        exceptions.append(
                            {
                                "sheet": sheet_name,
                                "row": row_index,
                                "reason": "手机号空或非法",
                            }
                        )
                        continue

                    name = row[column_index[name_col]] if name_col else ""
                    org_text = row[column_index[org_col]] if org_col else ""
                    region_text = row[column_index[region_col]] if region_col else ""
                    title_text = row[column_index[title_col]] if title_col else ""
                    remote_id_snapshot = (
                        row[column_index[remote_id_col]] if remote_id_col else ""
                    )
                    room_preference = row[column_index[room_col]] if room_col else ""

                    cursor = conn.execute(
                        "SELECT person_id FROM person WHERE phone_norm = ?",
                        (phone_norm,),
                    )
                    person_row = cursor.fetchone()
                    if person_row:
                        person_id = person_row["person_id"]
                        conn.execute(
                            """
                            UPDATE person
                            SET name_latest = ?, org_text_latest = ?
                            WHERE person_id = ?
                            """,
                            (name or None, org_text or None, person_id),
                        )
                    else:
                        cursor = conn.execute(
                            """
                            INSERT INTO person (phone_norm, name_latest, org_text_latest)
                            VALUES (?, ?, ?)
                            """,
                            (phone_norm, name or None, org_text or None),
                        )
                        person_id = cursor.lastrowid
                        new_person_count += 1

                    conn.execute(
                        """
                        INSERT INTO enrollment (
                            session_id,
                            person_id,
                            enrolled_at,
                            name_snapshot,
                            org_text,
                            region_text,
                            title_text,
                            remote_id_snapshot,
                            room_preference,
                            source_file,
                            source_sheet
                        )
                        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                        """,
                        (
                            session_id,
                            person_id,
                            datetime.now().isoformat(timespec="seconds"),
                            name or None,
                            org_text or None,
                            region_text or None,
                            title_text or None,
                            remote_id_snapshot or None,
                            room_preference or None,
                            source_file,
                            sheet_name,
                        ),
                    )
                    new_enrollment_count += 1
                    valid_rows += 1
            conn.commit()
        except Exception:
            conn.rollback()
            raise

    return {
        "sheet_count": sheet_count,
        "valid_rows": valid_rows,
        "new_person_count": new_person_count,
        "new_enrollment_count": new_enrollment_count,
        "exceptions": exceptions,
    }


@app.route("/api/enrollment/import", methods=["POST"])
def import_enrollment():
    excel_file = request.files.get("excel_file")
    if not excel_file or not excel_file.filename:
        return json_response(False, error="请上传报名 Excel 文件。")
    if not is_excel_filename(excel_file.filename):
        return json_response(
            False,
            error="仅支持 Excel 文件（.xlsx/.xls/.xlsm/.xltx/.xltm），请重新上传。",
        )

    session_id_value = request.form.get("session_id")
    session_id = resolve_session_id(session_id_value)
    if not session_id:
        return json_response(False, error="未找到可用的期次，请先创建期次。")

    cursor = get_connection().execute(
        "SELECT session_id FROM training_session WHERE session_id = ?", (session_id,)
    )
    if not cursor.fetchone():
        return json_response(False, error="期次不存在，请重新创建。")

    source_file, file_path = save_upload(excel_file)
    try:
        receipt = import_excel(file_path, source_file, session_id)
    except Exception as exc:
        return json_response(False, error=f"导入失败: {exc}")

    return json_response(True, receipt)




def normalize_cell_text(text: str) -> str:
    return re.sub(r"\s+", " ", (text or "").strip())


def find_course_column_indexes(headers: List[str]) -> Optional[Dict[str, int]]:
    mapping: Dict[str, int] = {}
    for idx, header in enumerate(headers):
        h = normalize_cell_text(header)
        if any(key in h for key in ["日期", "日 期"]):
            mapping["date"] = idx
        elif any(key in h for key in ["时间", "时 间"]):
            mapping["time"] = idx
        elif any(key in h for key in ["内容", "课程", "课程名称"]):
            mapping["course"] = idx
        elif any(key in h for key in ["授课教师", "教师", "老师"]):
            mapping["teacher"] = idx

    if "course" not in mapping:
        return None
    return mapping


def parse_date_text(date_text: str, default_year: int) -> Optional[date]:
    text = normalize_cell_text(date_text)
    if not text:
        return None
    nums = re.findall(r"\d+", text)
    if len(nums) >= 3:
        year, month, day = int(nums[0]), int(nums[1]), int(nums[2])
    elif len(nums) >= 2:
        year, month, day = default_year, int(nums[0]), int(nums[1])
    else:
        return None
    try:
        return date(year, month, day)
    except ValueError:
        return None


def parse_time_range(time_text: str) -> Tuple[Optional[datetime], Optional[datetime]]:
    text = normalize_cell_text(time_text).replace("：", ":")
    times = re.findall(r"(\d{1,2}:\d{2})", text)
    if len(times) >= 2:
        return times[0], times[1]
    if len(times) == 1:
        return times[0], None
    return None, None


def combine_date_time(day: Optional[date], time_val: Optional[str]) -> Optional[str]:
    if not day:
        return None
    if not time_val:
        return datetime(day.year, day.month, day.day, 0, 0).isoformat(timespec="seconds")
    hour, minute = [int(x) for x in time_val.split(":", 1)]
    return datetime(day.year, day.month, day.day, hour, minute).isoformat(timespec="seconds")


def parse_course_rows_from_word(
    file_path: str, default_year: int, location_text: str = "", session_id: Optional[int] = None
) -> List[Dict[str, Any]]:
    document = Document(file_path)
    records: List[Dict[str, Any]] = []

    for table in document.tables:
        rows = [
            [normalize_cell_text(cell.text) for cell in row.cells]
            for row in table.rows
        ]
        if len(rows) < 2:
            continue

        mapping = find_course_column_indexes(rows[0])
        if not mapping:
            continue

        last_date_text = ""
        last_time_text = ""
        for row in rows[1:]:
            course_name = row[mapping["course"]].strip() if mapping["course"] < len(row) else ""
            if not course_name or course_name in {"报到", "返程", "下午", "上午"}:
                continue

            date_text = row[mapping["date"]].strip() if mapping.get("date") is not None and mapping["date"] < len(row) else ""
            time_text = row[mapping["time"]].strip() if mapping.get("time") is not None and mapping["time"] < len(row) else ""
            teacher_name = row[mapping["teacher"]].strip() if mapping.get("teacher") is not None and mapping["teacher"] < len(row) else ""

            if date_text:
                last_date_text = date_text
            if time_text:
                last_time_text = time_text

            final_date_text = date_text or last_date_text
            final_time_text = time_text or last_time_text

            day = parse_date_text(final_date_text, default_year)
            start_time, end_time = parse_time_range(final_time_text)
            start_at = combine_date_time(day, start_time)
            end_at = combine_date_time(day, end_time)

            records.append(
                {
                    "title": course_name,
                    "teacher": teacher_name,
                    "start_at": start_at,
                    "end_at": end_at,
                    "location": location_text,
                    "session_id": session_id,
                }
            )

    return records


def build_qr_data_uri(text: str) -> str:
    if not QR_PIL_AVAILABLE:
        raise RuntimeError("未安装 Pillow（PIL），无法生成二维码。请先安装 qrcode[pil]。")
    image = qrcode.make(text)
    buffer = io.BytesIO()
    image.save(buffer, format="PNG")
    encoded = base64.b64encode(buffer.getvalue()).decode("utf-8")
    return f"data:image/png;base64,{encoded}"


def create_today_tasks() -> Dict[str, int]:
    today = date.today().isoformat()
    generated = 0
    skipped = 0
    qr_warning_logged = False

    with get_connection() as conn:
        courses = conn.execute(
            """
            SELECT course_id, title, teacher, start_at, end_at, location
            FROM course
            WHERE substr(start_at, 1, 10) = ?
            """,
            (today,),
        ).fetchall()

        for course in courses:
            course_id = course["course_id"]
            title = course["title"] or "课程"
            teacher = course["teacher"] or ""
            location = course["location"] or ""
            start_at = datetime.fromisoformat(course["start_at"]) if course["start_at"] else None
            end_at = datetime.fromisoformat(course["end_at"]) if course["end_at"] else None

            post_planned = end_at or start_at

            tasks = [
                ("post", post_planned, f"【课后问卷】请填写 {title} 的反馈问卷。"),
            ]

            for task_type, planned, content in tasks:
                if not planned:
                    skipped += 1
                    continue
                planned_iso = planned.isoformat(timespec="seconds")
                survey_link = None
                qr_data_uri = None
                if task_type == "post":
                    survey_link = f"http://127.0.0.1:5000/survey/{course_id}"
                    if not QR_PIL_AVAILABLE:
                        qr_data_uri = None
                        if not qr_warning_logged:
                            app.logger.error(
                                "QR generation disabled: Pillow(PIL) missing. Install with: pip install qrcode[pil]"
                            )
                            qr_warning_logged = True
                    else:
                        try:
                            qr_data_uri = build_qr_data_uri(survey_link)
                        except Exception as exc:
                            qr_data_uri = None
                            app.logger.exception("QR generation failed for course_id=%s: %s", course_id, exc)

                try:
                    conn.execute(
                        """
                        INSERT INTO message_task (
                            course_id, task_type, planned_at, content, survey_link,
                            qr_data_uri, status, created_at
                        )
                        VALUES (?, ?, ?, ?, ?, ?, 'pending', ?)
                        """,
                        (
                            course_id,
                            task_type,
                            planned_iso,
                            content,
                            survey_link,
                            qr_data_uri,
                            datetime.now().isoformat(timespec="seconds"),
                        ),
                    )
                    generated += 1
                except sqlite3.IntegrityError:
                    skipped += 1
                    continue
        conn.commit()

    return {"generated": generated, "skipped": skipped}


@app.route("/api/course/import", methods=["POST"], endpoint="api_course_import")
def import_course_word():
    word_file = request.files.get("word_file")
    if not word_file or not word_file.filename:
        return json_response(False, error="请上传 Word 课程表文件。")
    if not is_word_filename(word_file.filename):
        return json_response(False, error="仅支持 .docx Word 文件。")

    session_id: Optional[int] = None
    default_year = date.today().year
    location_text = ""
    session_id_text = request.form.get("session_id", "").strip()
    if session_id_text:
        if not session_id_text.isdigit():
            return json_response(False, error="session_id 非法。")
        session_id = int(session_id_text)
        with get_connection() as conn:
            session_row = conn.execute(
                "SELECT session_id, start_date, location_text FROM training_session WHERE session_id = ?",
                (session_id,),
            ).fetchone()
        if not session_row:
            return json_response(False, error="绑定的期次不存在。")
        location_text = session_row["location_text"] or ""
        start_date = session_row["start_date"] or ""
        if re.fullmatch(r"\d{4}-\d{2}-\d{2}", start_date):
            default_year = int(start_date[:4])

    source_file, file_path = save_upload(word_file)
    try:
        rows = parse_course_rows_from_word(file_path, default_year, location_text, session_id)
    except Exception as exc:
        return json_response(False, error=f"Word 读取失败: {exc}")

    if not rows:
        return json_response(False, error="未在 Word 表格中识别到课程内容列，请检查表头。")

    with get_connection() as conn:
        try:
            for row in rows:
                conn.execute(
                    """
                    INSERT INTO course (
                        title, teacher, start_at, end_at, location, session_id,
                        source_file, created_at
                    )
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?)
                    """,
                    (
                        row["title"],
                        row["teacher"] or None,
                        row["start_at"],
                        row["end_at"],
                        row["location"] or None,
                        row["session_id"],
                        source_file,
                        datetime.now().isoformat(timespec="seconds"),
                    ),
                )
            conn.commit()
        except Exception:
            conn.rollback()
            raise

    return json_response(True, {"imported_courses": len(rows), "rows": rows[:20]})


@app.route("/api/course/list")
def list_courses():
    with get_connection() as conn:
        rows = conn.execute(
            """
            SELECT course_id, title, teacher, start_at, end_at, location,
                   session_id, source_file, created_at
            FROM course
            ORDER BY COALESCE(start_at, created_at) DESC, course_id DESC
            LIMIT 300
            """
        ).fetchall()
    return json_response(True, [dict(row) for row in rows])


@app.route("/api/course/create", methods=["POST"])
def create_course():
    payload = request.get_json(silent=True) or {}
    title = str(payload.get("title", "")).strip()
    if not title:
        return json_response(False, error="课程名称不能为空。")

    with get_connection() as conn:
        cursor = conn.execute(
            """
            INSERT INTO course (title, teacher, start_at, end_at, location, session_id, source_file, created_at)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                title,
                (payload.get("teacher") or "").strip() or None,
                (payload.get("start_at") or "").strip() or None,
                (payload.get("end_at") or "").strip() or None,
                (payload.get("location") or "").strip() or None,
                int(payload["session_id"]) if str(payload.get("session_id", "")).isdigit() else None,
                "manual",
                datetime.now().isoformat(timespec="seconds"),
            ),
        )
    return json_response(True, {"course_id": cursor.lastrowid})


@app.route("/api/course/<int:course_id>")
def get_course(course_id: int):
    with get_connection() as conn:
        row = conn.execute(
            """
            SELECT course_id, title, teacher, start_at, end_at, location, session_id, source_file, created_at
            FROM course WHERE course_id = ?
            """,
            (course_id,),
        ).fetchone()
    if not row:
        return json_response(False, error="课程不存在。")
    return json_response(True, dict(row))


@app.route("/api/course/update", methods=["POST"])
def update_course():
    payload = request.get_json(silent=True) or {}
    course_id = payload.get("course_id")
    if not isinstance(course_id, int):
        return json_response(False, error="course_id 非法。")

    title = str(payload.get("title", "")).strip()
    if not title:
        return json_response(False, error="课程名称不能为空。")

    with get_connection() as conn:
        row = conn.execute("SELECT course_id FROM course WHERE course_id = ?", (course_id,)).fetchone()
        if not row:
            return json_response(False, error="课程不存在。")

        conn.execute(
            """
            UPDATE course
            SET title = ?, teacher = ?, start_at = ?, end_at = ?, location = ?, session_id = ?
            WHERE course_id = ?
            """,
            (
                title,
                (payload.get("teacher") or "").strip() or None,
                (payload.get("start_at") or "").strip() or None,
                (payload.get("end_at") or "").strip() or None,
                (payload.get("location") or "").strip() or None,
                int(payload["session_id"]) if str(payload.get("session_id", "")).isdigit() else None,
                course_id,
            ),
        )
    return json_response(True, {"course_id": course_id})


@app.route("/api/course/<int:course_id>/delete", methods=["POST"])
def delete_course(course_id: int):
    with get_connection() as conn:
        conn.execute("DELETE FROM message_task WHERE course_id = ?", (course_id,))
        conn.execute("DELETE FROM survey_response WHERE course_id = ?", (course_id,))
        conn.execute("DELETE FROM course WHERE course_id = ?", (course_id,))
    return json_response(True, {"course_id": course_id})


@app.route("/api/tasks/generate_today", methods=["POST"])
def generate_today_tasks():
    result = create_today_tasks()
    return json_response(True, result)


@app.route("/api/tasks/today")
def list_today_tasks():
    amap_key = request.args.get("map_api_key", "").strip()
    with get_connection() as conn:
        rows = conn.execute(
            """
            SELECT c.course_id,
                   c.title AS course_title,
                   c.teacher AS course_teacher,
                   c.location AS course_location,
                   c.start_at,
                   c.end_at,
                   mt.task_id,
                   mt.task_type,
                   mt.planned_at,
                   mt.content,
                   mt.survey_link,
                   mt.qr_data_uri,
                   mt.status,
                   mt.sent_at,
                   (
                       SELECT COUNT(1)
                       FROM survey_response sr
                       WHERE sr.course_id = c.course_id
                   ) AS survey_submitted_count,
                   (
                       SELECT COUNT(1)
                       FROM enrollment e
                       WHERE e.session_id = c.session_id
                   ) AS enrollment_total_count
            FROM course c
            LEFT JOIN message_task mt
              ON mt.course_id = c.course_id AND mt.task_type = 'post'
            ORDER BY COALESCE(c.end_at, c.start_at, c.created_at) DESC, c.course_id DESC
            LIMIT 3
            """,
        ).fetchall()
    result = []
    for row in rows:
        item = dict(row)
        if not item.get("content"):
            item["content"] = f"【课后问卷】请填写 {item.get('course_title') or '课程'} 的反馈问卷。"
        if not item.get("planned_at"):
            item["planned_at"] = item.get("end_at") or item.get("start_at") or ""
        if not item.get("status"):
            item["status"] = "pending"
        if not item.get("survey_link"):
            item["survey_link"] = f"http://127.0.0.1:5000/survey/{item['course_id']}"
        map_info = build_map_info(item.get("course_location", ""), amap_key)
        item["map_url"] = map_info["map_url"]
        item["geo"] = map_info["geo"]
        result.append(item)
    return json_response(True, result)


@app.route("/api/tasks/upsert_post", methods=["POST"])
def upsert_post_task():
    payload = request.get_json(silent=True) or {}
    course_id = payload.get("course_id")
    if not isinstance(course_id, int):
        return json_response(False, error="course_id 非法。")

    content = str(payload.get("content", "")).strip()
    if not content:
        return json_response(False, error="发送内容不能为空。")

    with get_connection() as conn:
        course = conn.execute(
            "SELECT course_id, start_at, end_at FROM course WHERE course_id = ?",
            (course_id,),
        ).fetchone()
        if not course:
            return json_response(False, error="课程不存在。")

        planned_at = course["end_at"] or course["start_at"] or datetime.now().isoformat(timespec="seconds")
        survey_link = f"http://127.0.0.1:5000/survey/{course_id}"
        existing = conn.execute(
            "SELECT task_id FROM message_task WHERE course_id = ? AND task_type = 'post' ORDER BY task_id DESC LIMIT 1",
            (course_id,),
        ).fetchone()
        if existing:
            conn.execute(
                """
                UPDATE message_task
                SET content = ?, planned_at = ?, survey_link = ?, status = 'pending', sent_at = NULL
                WHERE task_id = ?
                """,
                (content, planned_at, survey_link, existing["task_id"]),
            )
            task_id = existing["task_id"]
        else:
            qr_data_uri = None
            if QR_PIL_AVAILABLE:
                try:
                    qr_data_uri = build_qr_data_uri(survey_link)
                except Exception:
                    qr_data_uri = None
            cursor = conn.execute(
                """
                INSERT INTO message_task (
                    course_id, task_type, planned_at, content, survey_link,
                    qr_data_uri, status, created_at
                ) VALUES (?, 'post', ?, ?, ?, ?, 'pending', ?)
                """,
                (
                    course_id,
                    planned_at,
                    content,
                    survey_link,
                    qr_data_uri,
                    datetime.now().isoformat(timespec="seconds"),
                ),
            )
            task_id = cursor.lastrowid

    return json_response(True, {"course_id": course_id, "task_id": task_id})


@app.route("/api/tasks/<int:task_id>/mark_sent", methods=["POST"])
def mark_task_sent(task_id: int):
    with get_connection() as conn:
        conn.execute(
            """
            UPDATE message_task
            SET status = 'sent', sent_at = ?
            WHERE task_id = ?
            """,
            (datetime.now().isoformat(timespec="seconds"), task_id),
        )
    return json_response(True, {"task_id": task_id})


@app.route("/api/logs/recent")
def recent_logs():
    lines = int(request.args.get("lines", "200"))
    lines = max(20, min(lines, 1000))
    if not LOG_PATH.exists():
        return json_response(True, {"log_path": str(LOG_PATH), "lines": []})

    content = LOG_PATH.read_text(encoding="utf-8", errors="replace").splitlines()
    return json_response(True, {"log_path": str(LOG_PATH), "lines": content[-lines:]})


@app.route("/survey/<int:course_id>")
def survey_page(course_id: int):
    with get_connection() as conn:
        course = conn.execute(
            "SELECT course_id, title, teacher FROM course WHERE course_id = ?",
            (course_id,),
        ).fetchone()
    if not course:
        return "课程不存在", 404
    return render_template("survey.html", course=dict(course))


@app.route("/api/survey/submit", methods=["POST"])
def submit_survey():
    payload = request.get_json(silent=True) or {}
    course_id = payload.get("course_id")
    if not isinstance(course_id, int):
        return json_response(False, error="course_id 非法。")

    with get_connection() as conn:
        course = conn.execute(
            "SELECT course_id FROM course WHERE course_id = ?",
            (course_id,),
        ).fetchone()
        if not course:
            return json_response(False, error="课程不存在。")

        conn.execute(
            """
            INSERT INTO survey_response (
                course_id, satisfaction_score, gain_text, suggestion_text,
                recommend_score, submitted_at
            ) VALUES (?, ?, ?, ?, ?, ?)
            """,
            (
                course_id,
                payload.get("satisfaction_score"),
                payload.get("gain_text"),
                payload.get("suggestion_text"),
                payload.get("recommend_score"),
                datetime.now().isoformat(timespec="seconds"),
            ),
        )
    return json_response(True, {"course_id": course_id})


def fetch_yearly_stats(year: str) -> Dict[str, Any]:
    with get_connection() as conn:
        enrollments = conn.execute(
            """
            SELECT enrollment.enrollment_id, person.phone_norm, person.name_latest
            FROM enrollment
            JOIN person ON enrollment.person_id = person.person_id
            JOIN training_session ON enrollment.session_id = training_session.session_id
            WHERE COALESCE(NULLIF(substr(training_session.start_date, 1, 4), ""), substr(enrollment.enrolled_at, 1, 4)) = ?
            """,
            (year,),
        ).fetchall()

    person_counts: Dict[str, Dict[str, Any]] = {}
    for row in enrollments:
        phone_norm = row["phone_norm"]
        if phone_norm not in person_counts:
            person_counts[phone_norm] = {
                "phone_norm": phone_norm,
                "name": row["name_latest"] or "",
                "count": 0,
            }
        person_counts[phone_norm]["count"] += 1

    total_enrollments = len(enrollments)
    total_people = len(person_counts)
    repeat_people = sum(1 for item in person_counts.values() if item["count"] >= 2)
    top5 = sorted(person_counts.values(), key=lambda x: x["count"], reverse=True)[:5]

    return {
        "total_enrollments": total_enrollments,
        "total_people": total_people,
        "repeat_people": repeat_people,
        "top5": top5,
    }


def normalize_header_name(text: str) -> str:
    text = (text or "").strip().replace("\ufeff", "")
    text = re.sub(r"^[0-9]+[\.、]", "", text)
    text = text.replace("（", "(").replace("）", ")")
    text = re.sub(r"\s+", "", text)
    return text.lower()


def choose_field(row: Dict[str, str], headers: Dict[str, str], aliases: List[str]) -> str:
    for alias in aliases:
        key = normalize_header_name(alias)
        if key in headers:
            return (row.get(headers[key], "") or "").strip()
    return ""


@app.route("/api/finance/import", methods=["POST"])
def import_finance_csv():
    csv_file = request.files.get("csv_file")
    if not csv_file or not csv_file.filename:
        return json_response(False, error="请上传财务 CSV 文件。")
    if not csv_file.filename.lower().endswith(".csv"):
        return json_response(False, error="仅支持 CSV 文件。")

    saved_name, file_path = save_upload(csv_file)
    imported = 0
    updated = 0
    skipped = 0

    try:
        with open(file_path, "r", encoding="utf-8-sig", errors="ignore", newline="") as handle:
            sample = handle.read(4096)
            handle.seek(0)
            try:
                dialect = csv.Sniffer().sniff(sample, delimiters=",\t;")
            except Exception:
                dialect = csv.excel_tab
            reader = csv.DictReader(handle, dialect=dialect)
            if not reader.fieldnames:
                return json_response(False, error="CSV 表头为空，无法导入。")

            headers = {normalize_header_name(name): name for name in reader.fieldnames}
            now = datetime.now().isoformat(timespec="seconds")
            with get_connection() as conn:
                for row in reader:
                    if not row:
                        continue
                    record_no = choose_field(row, headers, ["编号", "id", "序号"])
                    if not record_no:
                        skipped += 1
                        continue

                    payload = {
                        "record_no": record_no,
                        "start_time": choose_field(row, headers, ["开始答题时间", "开始时间"]),
                        "end_time": choose_field(row, headers, ["结束答题时间", "结束时间"]),
                        "duration_text": choose_field(row, headers, ["答题时长", "时长"]),
                        "name": choose_field(row, headers, ["姓名", "1.姓名"]),
                        "phone": choose_field(row, headers, ["手机", "手机号", "2.手机"]),
                        "id_card": choose_field(row, headers, ["身份证号", "3.身份证号"]),
                        "org_name": choose_field(row, headers, ["工作单位", "4.工作单位"]),
                        "job_title": choose_field(row, headers, ["职务/职称", "5.职务/职称", "职务"]),
                        "bank_card": choose_field(row, headers, ["银行卡号", "7.银行卡号"]),
                        "bank_name": choose_field(row, headers, ["开户行", "8.开户行"]),
                        "city_name": choose_field(row, headers, ["地理位置市", "城市"]),
                        "user_type": choose_field(row, headers, ["用户类型"]),
                        "nickname": choose_field(row, headers, ["昵称"]),
                        "source_file": saved_name,
                        "updated_at": now,
                        "raw_json": json.dumps(row, ensure_ascii=False),
                    }

                    exists = conn.execute(
                        "SELECT record_id FROM finance_record WHERE record_no = ?",
                        (record_no,),
                    ).fetchone()
                    if exists:
                        conn.execute(
                            """
                            UPDATE finance_record
                            SET start_time = ?, end_time = ?, duration_text = ?, name = ?, phone = ?,
                                id_card = ?, org_name = ?, job_title = ?, bank_card = ?, bank_name = ?,
                                city_name = ?, user_type = ?, nickname = ?, source_file = ?, updated_at = ?, raw_json = ?
                            WHERE record_no = ?
                            """,
                            (
                                payload["start_time"], payload["end_time"], payload["duration_text"], payload["name"], payload["phone"],
                                payload["id_card"], payload["org_name"], payload["job_title"], payload["bank_card"], payload["bank_name"],
                                payload["city_name"], payload["user_type"], payload["nickname"], payload["source_file"], payload["updated_at"], payload["raw_json"],
                                record_no,
                            ),
                        )
                        updated += 1
                    else:
                        conn.execute(
                            """
                            INSERT INTO finance_record (
                                record_no, start_time, end_time, duration_text, name, phone, id_card,
                                org_name, job_title, bank_card, bank_name, city_name, user_type,
                                nickname, source_file, updated_at, raw_json
                            ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                            """,
                            (
                                payload["record_no"], payload["start_time"], payload["end_time"], payload["duration_text"], payload["name"], payload["phone"], payload["id_card"],
                                payload["org_name"], payload["job_title"], payload["bank_card"], payload["bank_name"], payload["city_name"], payload["user_type"],
                                payload["nickname"], payload["source_file"], payload["updated_at"], payload["raw_json"],
                            ),
                        )
                        imported += 1
                conn.commit()
    except Exception as exc:
        app.logger.exception("Finance CSV import failed")
        return json_response(False, error=f"导入失败：{exc}")

    return json_response(True, {
        "imported": imported,
        "updated": updated,
        "skipped": skipped,
        "source_file": saved_name,
    })


@app.route("/api/finance/list")
def finance_list():
    keyword = request.args.get("q", "").strip()
    page = max(1, int(request.args.get("page", "1")))
    page_size = min(100, max(10, int(request.args.get("page_size", "20"))))
    offset = (page - 1) * page_size

    where_sql = ""
    params: List[Any] = []
    if keyword:
        where_sql = "WHERE record_no LIKE ? OR name LIKE ? OR phone LIKE ? OR org_name LIKE ? OR bank_name LIKE ?"
        like_kw = f"%{keyword}%"
        params.extend([like_kw, like_kw, like_kw, like_kw, like_kw])

    with get_connection() as conn:
        total = conn.execute(
            f"SELECT COUNT(1) AS c FROM finance_record {where_sql}",
            tuple(params),
        ).fetchone()["c"]
        rows = conn.execute(
            f"""
            SELECT record_id, record_no, start_time, end_time, duration_text, name, phone,
                   id_card, org_name, job_title, bank_card, bank_name, city_name,
                   user_type, nickname, source_file, updated_at
            FROM finance_record
            {where_sql}
            ORDER BY COALESCE(start_time, updated_at) DESC, record_id DESC
            LIMIT ? OFFSET ?
            """,
            tuple(params + [page_size, offset]),
        ).fetchall()

    return json_response(True, {
        "total": total,
        "page": page,
        "page_size": page_size,
        "rows": [dict(row) for row in rows],
    })


def split_teacher_names(text: str) -> List[str]:
    raw = (text or "").strip()
    if not raw:
        return []
    parts = re.split(r"[、,/，;；\s]+", raw)
    return [p.strip() for p in parts if p.strip()]


@app.route("/api/finance/export/session_teachers")
def export_session_teachers_finance():
    session_id_text = request.args.get("session_id", "").strip()
    if not session_id_text.isdigit():
        return json_response(False, error="请提供合法 session_id。")
    session_id = int(session_id_text)

    with get_connection() as conn:
        session_row = conn.execute(
            """
            SELECT session_id, title, start_date, end_date, location_text
            FROM training_session
            WHERE session_id = ?
            """,
            (session_id,),
        ).fetchone()
        if not session_row:
            return json_response(False, error="培训班不存在。")

        course_rows = conn.execute(
            """
            SELECT course_id, title, teacher, start_at, end_at, location
            FROM course
            WHERE session_id = ?
            ORDER BY COALESCE(start_at, end_at, created_at), course_id
            """,
            (session_id,),
        ).fetchall()

        teacher_map: Dict[str, Dict[str, Any]] = {}
        for course in course_rows:
            names = split_teacher_names(course["teacher"] or "")
            if not names:
                continue
            for teacher_name in names:
                if teacher_name not in teacher_map:
                    teacher_map[teacher_name] = {
                        "teacher_name": teacher_name,
                        "course_titles": set(),
                        "first_start_at": course["start_at"] or "",
                        "last_end_at": course["end_at"] or "",
                        "course_location": course["location"] or "",
                    }
                teacher_map[teacher_name]["course_titles"].add(course["title"] or "")
                if course["start_at"] and (
                    not teacher_map[teacher_name]["first_start_at"]
                    or course["start_at"] < teacher_map[teacher_name]["first_start_at"]
                ):
                    teacher_map[teacher_name]["first_start_at"] = course["start_at"]
                if course["end_at"] and (
                    not teacher_map[teacher_name]["last_end_at"]
                    or course["end_at"] > teacher_map[teacher_name]["last_end_at"]
                ):
                    teacher_map[teacher_name]["last_end_at"] = course["end_at"]

        export_rows: List[Dict[str, Any]] = []
        for teacher_name, info in teacher_map.items():
            finance = conn.execute(
                """
                SELECT name, phone, id_card, org_name, job_title, bank_card, bank_name, city_name, user_type, nickname, updated_at
                FROM finance_record
                WHERE name = ?
                ORDER BY COALESCE(start_time, updated_at) DESC, record_id DESC
                LIMIT 1
                """,
                (teacher_name,),
            ).fetchone()
            export_rows.append(
                {
                    "培训班ID": session_row["session_id"],
                    "培训班名称": session_row["title"] or "",
                    "培训时间": f"{session_row['start_date'] or ''}~{session_row['end_date'] or ''}",
                    "培训地点": session_row["location_text"] or "",
                    "授课老师": teacher_name,
                    "课程名称": "；".join(sorted(info["course_titles"])),
                    "课程开始": info["first_start_at"],
                    "课程结束": info["last_end_at"],
                    "课程地点": info["course_location"],
                    "手机": finance["phone"] if finance else "",
                    "身份证号": finance["id_card"] if finance else "",
                    "工作单位": finance["org_name"] if finance else "",
                    "职务/职称": finance["job_title"] if finance else "",
                    "银行卡号": finance["bank_card"] if finance else "",
                    "开户行": finance["bank_name"] if finance else "",
                    "城市": finance["city_name"] if finance else "",
                    "用户类型": finance["user_type"] if finance else "",
                    "昵称": finance["nickname"] if finance else "",
                    "财务信息更新时间": finance["updated_at"] if finance else "",
                }
            )

    if not export_rows:
        return json_response(False, error="该培训班课程表中未找到授课老师信息，无法导出。")

    df = pd.DataFrame(export_rows)
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine="openpyxl") as writer:
        df.to_excel(writer, index=False, sheet_name="老师财务信息")
    output.seek(0)
    filename = f"session_{session_id}_teachers_finance.xlsx"
    return send_file(
        output,
        as_attachment=True,
        download_name=filename,
        mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )




@app.route("/api/session/history")
def session_history():
    with get_connection() as conn:
        rows = conn.execute(
            """
            SELECT
                ts.session_id,
                ts.title,
                ts.start_date,
                ts.end_date,
                ts.location_text,
                ts.training_goal,
                ts.created_at,
                COUNT(e.enrollment_id) AS enrollment_count
            FROM training_session ts
            LEFT JOIN enrollment e ON e.session_id = ts.session_id
            GROUP BY ts.session_id
            ORDER BY ts.session_id DESC
            """
        ).fetchall()

    return json_response(True, [dict(row) for row in rows])

@app.route("/api/stats/year")
def stats_year():
    year = request.args.get("year", "").strip()
    if not re.fullmatch(r"\d{4}", year):
        return json_response(False, error="请输入四位年份。")

    stats = fetch_yearly_stats(year)
    return json_response(True, stats)


def build_exports(year: str) -> io.BytesIO:
    with get_connection() as conn:
        enrollments = conn.execute(
            """
            SELECT enrollment.*, person.phone_norm, person.name_latest, training_session.title AS session_title,
                   training_session.start_date, training_session.end_date, training_session.location_text
            FROM enrollment
            JOIN person ON enrollment.person_id = person.person_id
            JOIN training_session ON enrollment.session_id = training_session.session_id
            WHERE COALESCE(NULLIF(substr(training_session.start_date, 1, 4), ""), substr(enrollment.enrolled_at, 1, 4)) = ?
            """,
            (year,),
        ).fetchall()

    enrollment_rows = [dict(row) for row in enrollments]
    enrollment_df = pd.DataFrame(enrollment_rows)

    if enrollment_df.empty:
        summary_df = pd.DataFrame(columns=["phone_norm", "name", "count"])
    else:
        summary_df = (
            enrollment_df.groupby(["phone_norm", "name_latest"], dropna=False)
            .size()
            .reset_index(name="count")
            .rename(columns={"name_latest": "name"})
            .sort_values("count", ascending=False)
        )

    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, "w", zipfile.ZIP_DEFLATED) as zf:
        enrollment_csv = enrollment_df.to_csv(index=False)
        zf.writestr(f"{year}_enrollments.csv", enrollment_csv)
        summary_csv = summary_df.to_csv(index=False)
        zf.writestr(f"{year}_person_summary.csv", summary_csv)
    buffer.seek(0)
    return buffer


@app.route("/api/export/year")
def export_year():
    year = request.args.get("year", "").strip()
    if not re.fullmatch(r"\d{4}", year):
        return json_response(False, error="请输入四位年份。")

    buffer = build_exports(year)
    filename = f"{year}_exports.zip"
    return send_file(
        buffer,
        as_attachment=True,
        download_name=filename,
        mimetype="application/zip",
    )


if __name__ == "__main__":
    setup_logging()
    initialize_database()
    print("本地服务已启动，请访问 http://127.0.0.1:5000")
    app.run(host="127.0.0.1", port=5000)
